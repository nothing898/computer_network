"""Generate the Computer Networking practice report draft.

Run with the bundled Python that has python-docx available:

    /path/to/bundled/python tools/generate_report.py
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "report"
OUT_FILE = OUT_DIR / "计网-xx组.docx"


def add_heading(document: Document, text: str, level: int = 1) -> None:
    paragraph = document.add_heading(text, level=level)
    for run in paragraph.runs:
        run.font.name = "SimHei"


def add_paragraph(document: Document, text: str = ""):
    paragraph = document.add_paragraph(text)
    paragraph.paragraph_format.first_line_indent = Pt(21)
    paragraph.paragraph_format.line_spacing = 1.5
    return paragraph


def add_code_block(document: Document, title: str, code: str) -> None:
    add_heading(document, title, 3)
    for line in code.strip("\n").splitlines():
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.line_spacing = 1.0
        run = paragraph.add_run(line)
        run.font.name = "Consolas"
        run.font.size = Pt(8)


def extract_block(source: str, marker: str) -> str:
    start = source.index(marker)
    method_match = marker.startswith("    def ")
    if method_match:
        next_match = source.find("\n    def ", start + len(marker))
        class_match = source.find("\n\nclass ", start + len(marker))
        candidates = [index for index in [next_match, class_match] if index != -1]
        end = min(candidates) if candidates else min(start + 2200, len(source))
    else:
        if marker == "class FTPClientApp":
            end = source.find("\n    def _build_local_panel", start)
            if end == -1:
                end = min(start + 2200, len(source))
        else:
            end = min(start + 2200, len(source))
    return source[start:end]


def configure_styles(document: Document) -> None:
    styles = document.styles
    styles["Normal"].font.name = "SimSun"
    styles["Normal"].font.size = Pt(10.5)
    for style_name in ["Heading 1", "Heading 2", "Heading 3"]:
        styles[style_name].font.name = "SimHei"


def build_report() -> None:
    OUT_DIR.mkdir(exist_ok=True)
    document = Document()
    configure_styles(document)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("武汉大学计算机学院本科生实验报告")
    run.bold = True
    run.font.name = "SimHei"
    run.font.size = Pt(18)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("计算机网络实践")
    run.font.name = "SimHei"
    run.font.size = Pt(16)

    for label in [
        "专业名称：XXX",
        "课程名称：计算机网络课程设计",
        "团队名称：XX组",
        "指导教师：吕慧",
        "团队成员一：XXX（20XXXXXXXXX）",
        "团队成员二：XXX（20XXXXXXXXX）",
        "日期：二○二六年六月",
    ]:
        paragraph = document.add_paragraph(label)
        paragraph.paragraph_format.space_after = Pt(6)

    document.add_page_break()

    add_heading(document, "郑重声明", 1)
    add_paragraph(
        document,
        "本团队呈交的实验报告，是在指导老师的指导下，独立进行实验工作所取得的成果，"
        "所有数据、图片资料真实可靠。尽我所知，除文中已经注明引用的内容外，本实验报告"
        "不包含他人享有著作权的内容。对本实验报告做出贡献的其他个人和集体，均已在文中"
        "以明确的方式标明。本实验报告的知识产权归属于培养单位。",
    )
    add_paragraph(document, "团队成员签名：________________      日期：2026年6月")

    add_heading(document, "摘要", 1)
    add_paragraph(
        document,
        "本实验完成了一个基于 socket 编程的图形化 FTP 客户端。系统从创建 TCP socket、"
        "建立 FTP 控制连接开始，实现用户登录、远程目录浏览、文件上传、文件下载以及断点续传。"
        "实验设计遵循 FTP 协议的控制连接与数据连接分离机制，使用 EPSV/PASV 被动模式建立"
        "数据连接，并使用 REST、RETR、STOR、APPE、SIZE 等命令完成文件传输与续传控制。",
    )
    add_paragraph(document, "关键词：FTP；Socket；图形化客户端；断点续传")

    add_heading(document, "1 实验目的和意义", 1)
    add_heading(document, "1.1 实验目的", 2)
    add_paragraph(
        document,
        "本实验是使学生熟悉网络规划与设计的基本知识和方法、掌握网络系统软件与应用软件开发的方法，"
        "能将所学的操作系统、数据库、软件工程、计算机网络等方面的知识集成到一起，规划、安装、"
        "调试实际网络系统、开发实际软件系统。本实验是使学生掌握网络系统软件的开发方法、"
        "开发平台的使用、与实际数据库的集成方法。用 Python 完成 FTP 客户端系统程序。",
    )
    add_heading(document, "1.2 实验意义", 2)
    add_paragraph(
        document,
        "该实验是理论知识和动手能力的综合体现。通过本实验，掌握网络系统软件、网络应用软件的"
        "开发方法、开发平台的使用、与实际网络协议的集成方法。FTP 客户端需要同时处理控制命令、"
        "数据连接、文件读写、异常恢复和图形界面状态更新，有助于理解应用层协议在 TCP 之上的"
        "实际工作过程。",
    )

    add_heading(document, "2 实验设计", 1)
    add_heading(document, "2.1 概述", 2)
    add_paragraph(
        document,
        "本系统选择 FTP 客户端题目，采用 Python 语言开发。界面部分使用 tkinter 实现，协议部分"
        "由 RawFTPClient 类完成。程序不调用 Python 标准库中的 ftplib，而是直接通过 socket"
        "创建控制连接和数据连接，实现 FTP 协议的主要命令流程。客户端提供服务器登录、远程目录"
        "刷新、远程目录切换、本地目录选择、文件上传、文件下载和传输进度显示等功能。",
    )

    add_heading(document, "2.2 实验原理", 2)
    add_paragraph(
        document,
        "FTP 使用两个 TCP 连接完成一次会话：控制连接用于发送 USER、PASS、LIST、RETR、STOR 等"
        "命令并接收三位状态码响应；数据连接用于传输目录列表或文件内容。本系统采用被动模式，"
        "客户端先发送 EPSV 或 PASV 命令，由服务器返回数据端口，随后客户端主动连接该端口并"
        "执行数据传输命令。下载续传时，客户端检测本地文件大小并发送 REST offset，再发送 RETR"
        "从指定偏移继续接收；上传续传时，客户端通过 SIZE 查询远程文件长度，随后使用 APPE 追加"
        "剩余数据，若远程文件不存在则使用 STOR 新建文件。",
    )

    add_heading(document, "2.3 实验方案", 2)
    add_paragraph(document, "系统分为协议层、界面层和任务调度层三个部分。")
    table = document.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    headers = ["模块", "主要职责", "关键实现"]
    for index, header in enumerate(headers):
        table.rows[0].cells[index].text = header
    rows = [
        ("协议层", "维护 FTP 控制连接和数据连接", "socket、EPSV/PASV、LIST、REST、RETR、STOR、APPE"),
        ("界面层", "展示本地与远程文件列表，收集用户操作", "tkinter、Treeview、Progressbar"),
        ("任务调度层", "避免网络传输阻塞界面", "threading 后台线程、queue 事件回传"),
        ("异常处理", "显示连接失败、登录失败、传输失败等错误", "FTPError、messagebox、日志窗口"),
    ]
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cells[index].text = value

    add_paragraph(
        document,
        "程序启动后，用户输入服务器地址、端口、用户名和密码。连接成功后客户端进入二进制传输模式，"
        "随后获取远程当前目录和文件列表。用户双击远程目录可进入目录，单击 Parent 可返回上级目录。"
        "上传和下载均在后台线程中进行，界面通过队列接收进度事件并更新进度条。",
    )

    add_heading(document, "2.4 小组成员分工", 2)
    division = document.add_table(rows=1, cols=3)
    division.style = "Table Grid"
    for index, header in enumerate(["成员", "分工", "对应源码"]):
        division.rows[0].cells[index].text = header
    for row in [
        ("XXX", "FTP 协议层设计与 socket 命令实现", "RawFTPClient、FTPResponse、RemoteEntry"),
        ("XXX", "图形界面、文件列表、上传下载交互", "FTPClientApp"),
        ("XXX", "测试、实验报告整理与运行说明", "README.md、实验报告"),
    ]:
        cells = division.add_row().cells
        for index, value in enumerate(row):
            cells[index].text = value

    add_heading(document, "结论", 1)
    add_heading(document, "1. 程序主要界面及结果", 2)
    add_paragraph(
        document,
        "程序主界面分为连接信息区、本地文件区、远程文件区、传输进度区和日志区。连接信息区输入"
        "FTP 服务器地址、端口、用户名和密码；本地文件区显示当前本地目录，可选择本地文件上传；"
        "远程文件区显示服务器目录，可双击目录进入并选择文件下载。上传和下载过程中进度条显示"
        "已传输字节数，日志区记录连接、登录、刷新、上传和下载结果。",
    )
    add_paragraph(
        document,
        "经测试，客户端能够完成 FTP 登录、目录列表获取、目录切换、普通上传、普通下载、下载断点续传"
        "和上传断点续传。下载续传通过保留本地未完成文件并再次执行下载验证，上传续传通过服务器端"
        "保留部分文件后再次上传验证。",
    )

    add_heading(document, "2. 程序源程序", 2)
    source = (ROOT / "src" / "ftp_client.py").read_text(encoding="utf-8")
    snippets = {
        "控制连接和命令响应读取": "    def command",
        "被动模式数据连接": "    def _open_passive_data_socket",
        "下载断点续传": "    def download",
        "上传断点续传": "    def upload",
        "图形界面入口": "class FTPClientApp",
    }
    for title_text, marker in snippets.items():
        add_code_block(document, title_text, extract_block(source, marker))

    add_heading(document, "参考文献", 1)
    for item in [
        "[1] RFC 959, File Transfer Protocol.",
        "[2] RFC 3659, Extensions to FTP.",
        "[3] Python Software Foundation. Python socket documentation.",
        "[4] Python Software Foundation. tkinter documentation.",
    ]:
        document.add_paragraph(item)

    add_heading(document, "【结论】", 1)
    add_paragraph(
        document,
        "本实验完成了符合要求的图形化 FTP 客户端。通过直接使用 socket 实现 FTP 控制连接和数据连接，"
        "进一步理解了应用层协议与 TCP 传输服务之间的关系，也掌握了文件传输程序中断点续传、"
        "异常处理和界面异步更新的基本实现方法。",
    )
    add_heading(document, "【小结】", 1)
    add_paragraph(
        document,
        "实验过程中较关键的问题是 FTP 多行响应解析、被动模式端口解析以及传输线程与 GUI 主线程之间"
        "的状态同步。通过将协议逻辑封装在 RawFTPClient 中、将界面逻辑封装在 FTPClientApp 中，"
        "程序结构较清晰，后续可继续扩展删除、重命名、新建目录、主动模式和 TLS 加密连接等功能。",
    )

    document.save(OUT_FILE)


if __name__ == "__main__":
    build_report()
